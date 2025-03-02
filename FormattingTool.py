import customtkinter as tk
from tkinter import filedialog
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.enum.text import WD_TAB_LEADER
from docx.enum.text import WD_TAB_ALIGNMENT


win = tk.CTk()
win.geometry("600x600")
win.resizable(False, False)
win.title("FormattingTool")


# Function to browse for a file and submit it
def browse_file():
    global second_part
    global first_part
    try:
        # Open file dialog to select a .docx file
        filepath = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        
        if filepath:  # Check if a file was selected
            # Attempt to open the document
            doc = Document(filepath)
            print("Document opened successfully.")
            new_window(doc)
            win.withdraw()

            first_part = []  # List for paragraphs before the second page break
            second_part = []  # List for paragraphs after the second page break
            page_break_count = 0


            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if 'lastRenderedPageBreak' in run._element.xml:  
                        page_break_count += 1 
                    if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                        page_break_count += 1
                    if run.contains_page_break:
                        page_break_count += 1
                if page_break_count < 2:
                    first_part.append(paragraph)  # Append to first_part before 2nd page break
                else:
                    second_part.append(paragraph)  # Append to second_part after 2nd page break

            print(f"First part: {len(first_part)} paragraphs")
            print(f"Second part: {len(second_part)} paragraphs")
        
        else:
            print("No file selected.")
    except Exception as e:
        # Handle any error during the opening of the document
        print(f"An error occurred while opening the document: {e}")

def has_bold(paragraph):
    # Check if any run in the paragraph is bold
    for run in paragraph.runs:
        if run.bold:
            return True
    return False

def modify_toc(doc):
    # Change document parameters
    for paragraph in toc_paragraphs:      
        for run in paragraph.runs:
            run.font.name = entry.get()
            if run.bold:
                run.font.size = Pt(float(font_size)) + Pt(2)
                print(run.font.size)
            else:
                run.font.size = Pt(float(font_size))  # Apply font size

        paragraph.paragraph_format.line_spacing = float(entry3.get())
        paragraph.paragraph_format.space_before = Pt(float(entry7.get()))
        paragraph.paragraph_format.space_after = Pt(float(entry7.get()))

        # Set alignment
        if entry4.get() == "center":
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif entry4.get() == "left":
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif entry4.get() == "right":
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif entry4.get() == "justify":
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    for para in toc_paragraphs:
        if has_bold(para):  # Check if paragraph contains bold text
            para.paragraph_format.space_before = Pt(float(entry8.get()))
            para.paragraph_format.space_after = Pt(float(entry8.get()))

            # Set title alignment
            if entry9.get() == "center":
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif entry9.get() == "left":
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif entry9.get() == "right":
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif entry9.get() == "justify":
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Update margins for all sections
        for section in doc.sections:
            section.top_margin = Cm(float(entry5.get()))
            section.bottom_margin = Cm(float(entry5.get()))
            section.left_margin = Cm(float(entry6.get()))
            section.right_margin = Cm(float(entry6.get()))


def modify_para(doc):

    global font_size
    global bold_font
    global final_font_size
    font_size = float(entry2.get())
    current_font_size = 0.0
    final_font_size = 0.0


    # Loop through paragraphs to find the current font size
    for paragraph in second_part:
        for run in paragraph.runs:
            if not run.bold:
                # Get the current font size in points
                current_font_size = run.font.size / 12700 if run.font.size is not None else font_size
                break  # Break after getting the first non bold font size

    # Calculate final font size
    final_font_size = (font_size - current_font_size) * 12700
    print(f"Final font size: {final_font_size}")

    
    # Change document parameters
    for paragraph in second_part:      
        for run in paragraph.runs:
            run.font.name = entry.get()
            if run.bold:
                bold_font = run.font.size + final_font_size if run.font.size else bold_font == Pt(float(font_size))
                run.font.size = bold_font
                print(run.font.size)
            else:
                run.font.size = Pt(float(font_size))  # Apply font size

        paragraph.paragraph_format.line_spacing = float(entry3.get())
        paragraph.paragraph_format.space_before = Pt(float(entry7.get()))
        paragraph.paragraph_format.space_after = Pt(float(entry7.get()))

        # Set alignment
        if entry4.get() == "center":
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif entry4.get() == "left":
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif entry4.get() == "right":
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif entry4.get() == "justify":
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    for para in second_part:
        if has_bold(para):  # Check if paragraph contains bold text
            para.paragraph_format.space_before = Pt(float(entry8.get()))
            para.paragraph_format.space_after = Pt(float(entry8.get()))

            # Set title alignment
            if entry9.get() == "center":
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif entry9.get() == "left":
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif entry9.get() == "right":
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif entry9.get() == "justify":
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Update margins for all sections
        for section in doc.sections:
            section.top_margin = Cm(float(entry5.get()))
            section.bottom_margin = Cm(float(entry5.get()))
            section.left_margin = Cm(float(entry6.get()))
            section.right_margin = Cm(float(entry6.get()))

def toc(doc, second_part):
    # Find the first paragraph after the 3rd page break
    page_break_count = 0
    global toc_paragraphs  # Store TOC paragraphs
    toc_paragraphs = []


    for i, paragraph in enumerate(doc.paragraphs):
        for run in paragraph.runs:
            if 'lastRenderedPageBreak' in run._element.xml:  
                page_break_count += 1 
            if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                page_break_count += 1
            if run.contains_page_break:
                page_break_count += 1

        if page_break_count == 2:
            # Insert TOC before the next paragraph
            if i + 1 < len(doc.paragraphs):
                toc_paragraph = doc.paragraphs[i].insert_paragraph_before()
                toc_title = toc_paragraph.add_run('Sisukord')
                toc_title.bold = True
                toc_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                toc_paragraphs.append(toc_paragraph)

                # Add TOC entries
                titles = []
                for para in second_part:
                    for run in para.runs:
                        if run.bold:
                            titles.append(para.text)
                            break  # Move to the next paragraph once a bold title is found
                
                # Add each title to the TOC with a manually assigned page number
                para_count = 1
                page_number = 1

                for title in titles:
                    toc_entry = doc.paragraphs[i + para_count].insert_paragraph_before()  # Create a paragraph for each entry
                    toc_entry.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.DOTS)
                    toc_run = toc_entry.add_run(title)
                    toc_run.add_text("\t")
                    run = toc_entry.add_run(f" {page_number}")
                    page_number += 1
                    para_count += 1
                    toc_paragraphs.append(toc_entry)

                toc_entry.add_run().add_break(WD_BREAK.PAGE)
                print(f"toc paragraphs: {len(toc_paragraphs)} paragraphs")

            return

def save_doc(doc):
    #choose the location and file name
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    
    if file_path:  # Check if a file path is chosen
        # If the user selected a location and file name, save the document
        if my_var.get() == "on":
            toc(doc, second_part)

        modify_para(doc)
        modify_toc(doc)
        doc.save(file_path)  # Save the document to the chosen location
        print(f"Document saved successfully at {file_path}")
    else:
        print("Save operation cancelled.")

def new_window(doc):
    new_win = tk.CTkToplevel(win)
    new_win.title("New Main Window")
    new_win.geometry("600x600")

    global entry
    global entry2
    global entry3
    global entry4
    global entry5
    global entry6
    global entry7
    global entry8
    global entry9
    global my_var
    
    
    section_text = tk.CTkLabel(new_win, text="Font formatting", width=120, height=0, font=("Arial", 15), text_color="black")
    section_text.place(x=245, y=10)

    entry = tk.CTkEntry(new_win, width=120)
    text = tk.CTkLabel(new_win, text="Font", font=("Arial", 12), text_color="black")
    text.place(x=90, y=40)
    entry.place(x=245, y=40)

    entry2 = tk.CTkEntry(new_win, width=120)
    text2 = tk.CTkLabel(new_win, text="Font size", font=("Arial", 12), text_color="black")
    text2.place(x=90, y=70)
    entry2.place(x=245, y=70)

    section_text2 = tk.CTkLabel(new_win, text="Spacing", width=120, height=0, font=("Arial", 15), text_color="black")
    section_text2.place(x=245, y=105)

    entry3 = tk.CTkEntry(new_win, width=120)
    text3 = tk.CTkLabel(new_win, text="Line Spacing", font=("Arial", 12), text_color="black")
    text3.place(x=90, y=130)
    entry3.place(x=245, y=130)

    entry7 = tk.CTkEntry(new_win, width=120)
    text7 = tk.CTkLabel(new_win, text="Paragraph spacing", font=("Arial", 12), text_color="black")
    text7.place(x=90, y=160)
    entry7.place(x=245, y=160)

    entry8 = tk.CTkEntry(new_win, width=120)
    text8 = tk.CTkLabel(new_win, text="Title spacing", font=("Arial", 12), text_color="black")
    text8.place(x=90, y=190)
    entry8.place(x=245, y=190)

    section_text3 = tk.CTkLabel(new_win, text="Alignment", width=120, height=0, font=("Arial", 15), text_color="black")
    section_text3.place(x=245, y=225)

    entry4 = tk.CTkEntry(new_win, width=120)
    text4 = tk.CTkLabel(new_win, text="Text alignment", font=("Arial", 12), text_color="black")
    text4.place(x=90, y=250)
    entry4.place(x=245, y=250)

    entry9 = tk.CTkEntry(new_win, width=120)
    text9 = tk.CTkLabel(new_win, text="Title alignment", font=("Arial", 12), text_color="black")
    text9.place(x=90, y=280)
    entry9.place(x=245, y=280)

    section_text4 = tk.CTkLabel(new_win, text="Margins", width=120, height=0, font=("Arial", 15), text_color="black")
    section_text4.place(x=245, y=315)

    entry5 = tk.CTkEntry(new_win, width=120)
    text5 = tk.CTkLabel(new_win, text="Top, Bottom margin", font=("Arial", 12), text_color="black")
    text5.place(x=90, y=340)
    entry5.place(x=245, y=340)

    entry6 = tk.CTkEntry(new_win, width=120)
    text6 = tk.CTkLabel(new_win, text="Left, Right margin", font=("Arial", 12), text_color="black")
    text6.place(x=90, y=370)
    entry6.place(x=245, y=370)

    my_var = tk.StringVar(value="default")
    switch = tk.CTkSwitch(new_win, text="Generate TOC", variable=my_var, onvalue="on", offvalue="off")
    switch.place(x=245, y=400)

    save_button = tk.CTkButton(new_win, text="Save", command=lambda: save_doc(doc))
    save_button.place(x=235, y=430)


# Button to open file browser
browse_button = tk.CTkButton(win, text="Browse", command=browse_file)
browse_button.pack(padx=20, pady=20)



win.mainloop()